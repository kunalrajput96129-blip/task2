import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

DOC_PATH = r"c:\Users\Vicky\OneDrive\Desktop\task2\Digital_Services_Portal_Design_Doc.docx"
ASSETS_DIR = r"c:\Users\Vicky\OneDrive\Desktop\task2\assets"

# Color Palette Constants for Word Document Styling
COLOR_NAVY = RGBColor(15, 23, 42)      # #0f172a
COLOR_BLUE = RGBColor(37, 99, 235)     # #2563eb
COLOR_SLATE = RGBColor(71, 85, 105)    # #475569
COLOR_GREEN = RGBColor(22, 163, 74)    # #16a34a

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell padding."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.bold = True
    
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = COLOR_NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = COLOR_BLUE
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_SLATE
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p

def add_callout(doc, text, title="DESIGN RATIONALE & KEY TAKEAWAY"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "EFF6FF") # Light Blue Tint
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left Border Thick Blue
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="36" w:space="0" w:color="2563EB"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r_title = p.add_run(f"💡 {title}\n")
    r_title.bold = True
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = COLOR_BLUE

    r_text = p.add_run(text)
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = COLOR_NAVY
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def generate_doc():
    doc = Document()

    # Configure Margins (1 inch)
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # ---------------------------------------------------------------------------
    # COVER PAGE
    # ---------------------------------------------------------------------------
    p_cover_top = doc.add_paragraph()
    p_cover_top.paragraph_format.space_before = Pt(36)
    p_cover_top.paragraph_format.space_after = Pt(12)
    r_pre = p_cover_top.add_run("WEEK 2 EXECUTION DELIVERABLE: DESIGN & PROTOTYPING")
    r_pre.font.size = Pt(11)
    r_pre.font.bold = True
    r_pre.font.color.rgb = COLOR_BLUE

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(12)
    r_main_title = p_title.add_run("Designing a Responsive Web Prototype for Digital Services")
    r_main_title.font.size = Pt(26)
    r_main_title.font.bold = True
    r_main_title.font.color.rgb = COLOR_NAVY

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(24)
    r_sub = p_sub.add_run("Comprehensive Design Rationale, High-Fidelity Wireframes, Accessibility Compliance Matrix (WCAG 2.1 AAA), and Front-End Architecture Specification")
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = COLOR_SLATE

    # Horizontal Divider Line
    p_div = doc.add_paragraph()
    r_div = p_div.add_run("_________________________________________________________________________________")
    r_div.font.color.rgb = COLOR_BLUE
    p_div.paragraph_format.space_after = Pt(36)

    # Metadata Block Table
    tbl_meta = doc.add_table(rows=4, cols=2)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Project Portal Name:", "GovDirect Unified Public Digital Services Portal"),
        ("Design Methodology:", "Mobile-First Fluid Grid & Universal Accessibility (WCAG 2.1 AAA)"),
        ("Target Users:", "General Citizens, Commercial Enterprises, Senior Citizens, Screen Reader Users"),
        ("Document Status:", "Final Executable Prototype Design Documentation")
    ]
    for idx, (label, val) in enumerate(meta_data):
        row = tbl_meta.rows[idx]
        cell_l, cell_r = row.cells[0], row.cells[1]
        set_cell_margins(cell_l, top=60, bottom=60, left=100, right=100)
        set_cell_margins(cell_r, top=60, bottom=60, left=100, right=100)
        
        p_l = cell_l.paragraphs[0]
        r1 = p_l.add_run(label)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = COLOR_NAVY

        p_r = cell_r.paragraphs[0]
        r2 = p_r.add_run(val)
        r2.font.size = Pt(10)
        r2.font.color.rgb = COLOR_SLATE

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # EXECUTIVE SUMMARY
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "Executive Summary", level=1)
    p_exec = doc.add_paragraph()
    p_exec.paragraph_format.space_after = Pt(10)
    p_exec.add_run(
        "This document details the complete design, architectural choices, responsive grid mechanics, and accessibility framework "
        "for the GovDirect Digital Services Portal prototype. Modern e-governance demands digital service portals that are not only visually "
        "stunning and responsive across all device form factors (Desktop, Tablet, Mobile), but also fully inclusive and accessible for citizens with "
        "differing physical, visual, and cognitive abilities."
    )

    add_callout(
        doc,
        "The primary goal of the GovDirect portal is to eliminate administrative friction in public service delivery. By combining an intuitive "
        "search catalog, real-time application tracking, multi-step validation forms, and built-in accessibility toolbars (font scaling, high-contrast dark mode, text-to-speech preview), "
        "the prototype fulfills government digital service standards while achieving WCAG 2.1 AAA compliance.",
        "EXECUTIVE OBJECTIVE"
    )

    # ---------------------------------------------------------------------------
    # CHAPTER 1: CONCEPTUALIZATION & PORTAL ARCHITECTURE
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "1. Conceptualization & Portal Architecture", level=1)
    p_c1 = doc.add_paragraph()
    p_c1.add_run(
        "The conceptualization phase identified core digital components necessary to deliver a seamless online civic portal. "
        "A structured component breakdown ensures high modularity, ease of maintenance, and universal user familiarity."
    )

    add_styled_heading(doc, "1.1 Essential Portal UI Components", level=2)
    
    # Table of Core UI Components
    comp_rows = [
        ("Accessibility Toolbar", "Provides top-level controls for dynamic font scaling (80%-140%), dark contrast toggle, motion reduction, and text-to-speech audio simulation.", "role='region', aria-label='Accessibility Controls'. Live updates announce mode changes."),
        ("Header Navigation Bar", "Provides persistent brand emblem, sticky menu items, quick reference ID search button, and responsive mobile hamburger trigger.", "role='banner', aria-expanded='false' for hamburger menu toggle."),
        ("Dynamic Search & Filter Hero", "Allows instant service search with keyword pattern matching and quick-filter chips for popular services (License, Business, Grants).", "role='search', aria-label='Search Government Services' with real-time DOM filtering."),
        ("Filterable Services Catalog", "Grid layout displaying digital service cards with processing SLA, fee details, accessibility badges, and application modal triggers.", "role='tablist' for category filter tabs; aria-selected='true' on active category."),
        ("Multi-Step Application Modal", "Step-by-step application drawer (Applicant Info → Document Upload → Review & Digital Signature) with real-time field validation.", "role='dialog', aria-modal='true', aria-labelledby='modal-title', focus trapping."),
        ("Real-Time Application Tracker", "Interactive widget allowing citizens to input reference codes (e.g. GOV-2026-8942) and view visual progress timeline.", "role='status', aria-live='polite' updating screen reader users on step progression.")
    ]

    tbl_comp = doc.add_table(rows=len(comp_rows) + 1, cols=3)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["UI Component", "Functionality & User Role", "Accessibility & ARIA Implementation"]
    
    hdr_cells = tbl_comp.rows[0].cells
    for i, h_text in enumerate(headers):
        set_cell_background(hdr_cells[i], "0F172A")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h_text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)


    for row_idx, data in enumerate(comp_rows, start=1):
        row_cells = tbl_comp.rows[row_idx].cells
        bg_hex = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            set_cell_background(row_cells[col_idx], bg_hex)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = COLOR_NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ---------------------------------------------------------------------------
    # CHAPTER 2: WIREFRAMES, LAYOUTS & USER FLOW ANALYSIS
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "2. High-Fidelity Wireframes, Layouts & User Flow Analysis", level=1)
    p_c2 = doc.add_paragraph()
    p_c2.add_run(
        "Wireframing establishes the spatial visual hierarchy, grid geometry, and interaction flows before code implementation. "
        "The following high-fidelity diagrams demonstrate layout adaptation from desktop monitors down to mobile viewports."
    )

    # Image 1: Desktop Wireframe
    path_desktop = os.path.join(ASSETS_DIR, "wireframe_desktop.png")
    if os.path.exists(path_desktop):
        add_styled_heading(doc, "2.1 Desktop Layout Wireframe (1200px Grid)", level=2)
        doc.add_picture(path_desktop, width=Inches(6.2))
        p_cap1 = doc.add_paragraph()
        p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap1 = p_cap1.add_run("Figure 1: High-Fidelity Wireframe for Desktop Portal Layout with Sidebar Filters and 3-Column Service Grid.")
        r_cap1.font.italic = True
        r_cap1.font.size = Pt(9)
        r_cap1.font.color.rgb = COLOR_SLATE

        add_callout(
            doc,
            "On desktop viewports (1024px and above), the layout utilizes a 3-column service grid flanked by a left category filter sidebar. "
            "The top persistent header incorporates the accessibility toolbar without obscuring primary navigation options.",
            "DESKTOP LAYOUT RATIONALE"
        )

    # Image 2: Mobile Wireframe
    path_mobile = os.path.join(ASSETS_DIR, "wireframe_mobile.png")
    if os.path.exists(path_mobile):
        add_styled_heading(doc, "2.2 Mobile Responsive Wireframe (375px Breakpoint)", level=2)
        doc.add_picture(path_mobile, width=Inches(3.8))
        p_cap2 = doc.add_paragraph()
        p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap2 = p_cap2.add_run("Figure 2: Mobile Responsive Layout Wireframe featuring Single-Column Stack, Collapsible Drawer Menu, and Sticky Navigation.")
        r_cap2.font.italic = True
        r_cap2.font.size = Pt(9)
        r_cap2.font.color.rgb = COLOR_SLATE

        add_callout(
            doc,
            "Mobile viewports compress navigation into a single vertical column. Buttons and interactive targets are scaled to a minimum touch target size of 48px × 48px to accommodate thumb-zone navigation.",
            "MOBILE RESPONSIVE RATIONALE"
        )

    # Image 3: Form Wireframe
    path_form = os.path.join(ASSETS_DIR, "wireframe_form.png")
    if os.path.exists(path_form):
        add_styled_heading(doc, "2.3 Interactive Multi-Step Application Modal Wireframe", level=2)
        doc.add_picture(path_form, width=Inches(5.8))
        p_cap3 = doc.add_paragraph()
        p_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap3 = p_cap3.add_run("Figure 3: Multi-step Application Form Modal Layout featuring Stepper Tabs, Inline Validation Badges, and File Upload Zone.")
        r_cap3.font.italic = True
        r_cap3.font.size = Pt(9)
        r_cap3.font.color.rgb = COLOR_SLATE

    # Image 4: User Flow
    path_uf = os.path.join(ASSETS_DIR, "user_flow.png")
    if os.path.exists(path_uf):
        add_styled_heading(doc, "2.4 Citizen User Flow & Interaction Architecture", level=2)
        doc.add_picture(path_uf, width=Inches(6.2))
        p_cap4 = doc.add_paragraph()
        p_cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap4 = p_cap4.add_run("Figure 4: End-to-End Citizen User Journey Flowchart from Discovery to Real-Time Progress Tracking.")
        r_cap4.font.italic = True
        r_cap4.font.size = Pt(9)
        r_cap4.font.color.rgb = COLOR_SLATE

    # ---------------------------------------------------------------------------
    # CHAPTER 3: RESPONSIVE DESIGN ENGINEERING
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "3. Responsive Design Engineering", level=1)
    p_c3 = doc.add_paragraph()
    p_c3.add_run(
        "Responsive web engineering ensures that the portal adapts gracefully across diverse screen resolutions without requiring horizontal scrolling or broken layouts. "
        "The prototype utilizes modern CSS layout modules including CSS Grid, Flexbox, and CSS fluid functions."
    )

    add_styled_heading(doc, "3.1 Breakpoint Matrix & Layout Behaviors", level=2)

    # Breakpoint Table
    tbl_bp = doc.add_table(rows=4, cols=4)
    tbl_bp.alignment = WD_TABLE_ALIGNMENT.CENTER
    bp_headers = ["Device Category", "Viewport Range", "Layout Mechanics", "UI Adaptations"]
    for i, h in enumerate(bp_headers):
        set_cell_background(tbl_bp.rows[0].cells[i], "0F172A")
        set_cell_margins(tbl_bp.rows[0].cells[i], top=80, bottom=80, left=100, right=100)
        p = tbl_bp.rows[0].cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9.5)

    bp_rows = [
        ("Mobile Handheld", "320px - 767px", "Single-column CSS Grid (`grid-template-columns: 1fr`)", "Navigation links collapse into hamburger overlay; font size uses `clamp()`; sticky bottom accessibility bar."),
        ("Tablet Portrait/Landscape", "768px - 1023px", "2-Column CSS Grid (`repeat(auto-fit, minmax(280px, 1fr))`)", "Top header displays simplified search box; filter category buttons convert into horizontal scroll strip."),
        ("Desktop & Ultrawide", "1024px and above", "3-Column Fluid CSS Grid with Fixed Sidebar", "Full header navigation visible; sidebar category filters stay persistent; accessibility toolbar pinned to top.")
    ]

    for row_idx, data in enumerate(bp_rows, start=1):
        row_cells = tbl_bp.rows[row_idx].cells
        bg_hex = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            set_cell_background(row_cells[col_idx], bg_hex)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9)
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = COLOR_NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_styled_heading(doc, "3.2 Principles of Fluid Grids & Typography", level=2)
    p_fluid = doc.add_paragraph()
    p_fluid.add_run(
        "Instead of hardcoded pixel font sizes, the prototype leverages CSS `clamp()` fluid functions to automatically interpolate text scale relative to viewport width:\n\n"
        "   `font-size: clamp(1.75rem, 3vw + 1rem, 2.75rem);`\n\n"
        "This formula prevents text overflow on mobile screens while ensuring high readability on large desktop monitors. "
        "Furthermore, images and SVG icons use flexible width declarations (`max-width: 100%; height: auto;`) to prevent layout blowout."
    )

    # ---------------------------------------------------------------------------
    # CHAPTER 4: USABILITY & INTERACTION DESIGN
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "4. Usability & Interaction Design", level=1)
    p_c4 = doc.add_paragraph()
    p_c4.add_run(
        "User-centered design principles guide the portal's interaction mechanics. Government forms are often criticized for high cognitive burden; "
        "the GovDirect prototype addresses this through structured progressive disclosure and instant feedback."
    )

    add_styled_heading(doc, "4.1 Key Usability Heuristics Implemented", level=2)
    usability_points = [
        ("Cognitive Load Reduction:", "Complex government applications are decomposed into manageable 3-step modal panels (Applicant Details → Document Upload → Review & Final Signoff)."),
        ("Error Prevention & Instant Recovery:", "Inline JavaScript validation verifies email formatting, phone numbers, and citizen ID inputs immediately on step transition, preventing frustrating post-submission rejections."),
        ("Visibility of System Status:", "The Live Tracker widget provides citizens with transparent visual step-by-step progress timelines for submitted reference IDs (GOV-2026-8942)."),
        ("Flexibility & Speed of Use:", "Frequent actions are accessible via quick-filter chips on the hero search bar, allowing 1-click access to Driver License, Business, and Grant services.")
    ]

    for title, desc in usability_points:
        p = doc.add_paragraph(style='List Bullet')
        r_t = p.add_run(title + " ")
        r_t.bold = True
        r_t.font.color.rgb = COLOR_NAVY
        r_d = p.add_run(desc)
        r_d.font.size = Pt(10)

    # ---------------------------------------------------------------------------
    # CHAPTER 5: ACCESSIBILITY & UNIVERSAL DESIGN MATRIX
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "5. Accessibility & Universal Design Matrix (WCAG 2.1 AAA)", level=1)
    p_c5 = doc.add_paragraph()
    p_c5.add_run(
        "Universal accessibility is a non-negotiable legal and moral mandate for government platforms. "
        "The prototype adheres strictly to Web Content Accessibility Guidelines (WCAG 2.1 AAA and AA levels)."
    )

    path_cm = os.path.join(ASSETS_DIR, "contrast_matrix.png")
    if os.path.exists(path_cm):
        add_styled_heading(doc, "5.1 Color Contrast Evaluation", level=2)
        doc.add_picture(path_cm, width=Inches(6.0))
        p_cap5 = doc.add_paragraph()
        p_cap5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap5 = p_cap5.add_run("Figure 5: Color Contrast Ratio Evaluation Matrix confirming WCAG 2.1 AAA/AA Compliance.")
        r_cap5.font.italic = True
        r_cap5.font.size = Pt(9)

    add_styled_heading(doc, "5.2 Universal Accessibility Features Breakdown", level=2)
    acc_features = [
        ("Dynamic Font Scaling Toolbar:", "Allows citizens with low vision to scale text from 80% up to 140% without breaking container dimensions or overflowing buttons."),
        ("High Contrast Dark Mode:", "Provides a 19.8:1 text contrast ratio option, reducing eye strain and aiding users with visual sensitivity."),
        ("Screen Reader Simulation & ARIA Live Regions:", "Integrated Web Speech API text-to-speech engine coupled with `aria-live='polite'` regions that announce live form errors and reference code creations."),
        ("Keyboard Focus Rings:", "All interactive buttons and input fields exhibit a prominent 3px solid focus ring (`outline: 3px solid #2563eb`), ensuring complete keyboard navigability (Tab/Shift+Tab).")
    ]

    for title, desc in acc_features:
        p = doc.add_paragraph(style='List Bullet')
        r_t = p.add_run(title + " ")
        r_t.bold = True
        r_t.font.color.rgb = COLOR_BLUE
        r_d = p.add_run(desc)
        r_d.font.size = Pt(10)

    # ---------------------------------------------------------------------------
    # CHAPTER 6: PROTOTYPE SOURCE CODE & EXECUTION GUIDE
    # ---------------------------------------------------------------------------
    add_styled_heading(doc, "6. Prototype Source Code & Execution Guide", level=1)
    p_c6 = doc.add_paragraph()
    p_c6.add_run(
        "The prototype is constructed as a self-contained front-end application using semantic HTML5, modern vanilla CSS custom properties, "
        "and client-side JavaScript. Below are execution instructions to launch and test the web prototype."
    )

    add_styled_heading(doc, "6.1 Directory & File Structure", level=2)
    p_tree = doc.add_paragraph()
    p_tree.paragraph_format.space_before = Pt(4)
    p_tree.paragraph_format.space_after = Pt(8)
    r_tree = p_tree.add_run(
        "task2/\n"
        "├── Digital_Services_Portal_Design_Doc.docx  (Comprehensive Design Document)\n"
        "├── generate_visual_wireframes.py            (Matplotlib/Pillow Diagram Generator)\n"
        "├── generate_doc_report.py                  (Python-Docx Report Generator)\n"
        "├── build_all.py                            (Master Build Orchestrator)\n"
        "├── assets/                                 (Generated Visual PNG Wireframes)\n"
        "└── prototype/                              (Web Prototype Source Code)\n"
        "    ├── index.html                          (Semantic HTML5 Markup)\n"
        "    ├── styles.css                          (CSS Design System & Breakpoints)\n"
        "    └── app.js                              (JavaScript Interactivity & ARIA Logic)\n"
    )
    r_tree.font.name = 'Courier New'
    r_tree.font.size = Pt(9.5)
    r_tree.font.color.rgb = COLOR_NAVY

    add_styled_heading(doc, "6.2 Steps to Run & Inspect Prototype", level=2)
    steps = [
        "1. Open terminal and navigate to directory: `cd c:\\Users\\Vicky\\OneDrive\\Desktop\\task2`",
        "2. To launch a local web server, run: `python -m http.server 8000 --directory prototype`",
        "3. Open web browser and visit: `http://localhost:8000`",
        "4. Interact with accessibility buttons (Font +/-, Dark Contrast Mode, Screen Reader TTS).",
        "5. Click 'Apply Now' on any service card to test the 3-step modal form and receive a generated tracking reference code.",
        "6. Enter the tracking reference code into the Live Application Tracker widget to observe progress updates."
    ]

    for s in steps:
        p = doc.add_paragraph()
        r = p.add_run(s)
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_SLATE

    # Save document
    doc.save(DOC_PATH)
    print(f"Generated Word Document: {DOC_PATH}")

if __name__ == "__main__":
    generate_doc()
